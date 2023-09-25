import openai
from docx import Document
from pydub import AudioSegment

#                   0           1          2                       3
option = int(input("transcribe, translate, transcribe + translate, all"))

def transcribe_audio(audio_file_path):
    return "===transcription==="
    with open(audio_file_path, 'rb') as audio_file:
        transcription = openai.Audio.transcribe("whisper-1", audio_file)
    return transcription['text']

def translation_chinese_extraction(transcription):
    return "===translation==="
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and translate it into chinese. Aim to retain the most important points, providing a coherent and readable translation that could help a person who speaks chinese understand the full text, however it is written in english. Please throroughly trsnalte the entire text, not mising out on any of the points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def abstract_summary_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a highly skilled AI trained in language comprehension and summarization. I would like you to read the following text and summarize it into a concise abstract paragraph. Aim to retain the most important points, providing a coherent and readable summary that could help a person understand the main points of the discussion without needing to read the entire text. Please avoid unnecessary details or tangential points."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def key_points_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def key_points_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are a proficient AI with a specialty in distilling information into key points. Based on the following text, identify and list the main points that were discussed or brought up. These should be the most important ideas, findings, or topics that are crucial to the essence of the discussion. Your goal is to provide a list that someone could read to quickly understand what was talked about."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']


def action_item_extraction(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "You are an AI expert in analyzing conversations and extracting action items. Please review the text and identify any tasks, assignments, or actions that were agreed upon or mentioned as needing to be done. These could be tasks assigned to specific individuals, or general actions that the group has decided to take. Please list these action items clearly and concisely."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def sentiment_analysis(transcription):
    response = openai.ChatCompletion.create(
        model="gpt-4",
        temperature=0,
        messages=[
            {
                "role": "system",
                "content": "As an AI with expertise in language and emotion analysis, your task is to analyze the sentiment of the following text. Please consider the overall tone of the discussion, the emotion conveyed by the language used, and the context in which words and phrases are used. Indicate whether the sentiment is generally positive, negative, or neutral, and provide brief explanations for your analysis where possible."
            },
            {
                "role": "user",
                "content": transcription
            }
        ]
    )
    return response['choices'][0]['message']['content']

def save_as_docx(minutes, filename):
    doc = Document()
    for key, value in minutes.items():
        # Replace underscores with spaces and capitalize each word for the heading
        heading = ' '.join(word.capitalize() for word in key.split('_'))
        doc.add_heading(heading, level=1)
        doc.add_paragraph(value)
        # Add a line break between sections
        doc.add_paragraph()
    doc.save(filename)

#transcription#################################################################################################################
def transcription(audio_path):
    audio = AudioSegment.from_mp3(audio_path)
    audio_len = len(audio)
    
    # PyDub handles time in milliseconds
    one_minute = 1 * 60 * 1000
    index = 0

    tmp_file_path = "foo.mp3"
    total_transcription = ""

    while audio_len > index * one_minute:
        minute = audio[index * one_minute : (index+1) * one_minute]
        minute.export(tmp_file_path, format="mp3")
        total_transcription += transcribe_audio(tmp_file_path)

        index += 1

    return(total_transcription)

#translation#################################################################################################################
def translation(transcript_path):
    if transcript_path != "empty":
        f = open(transcript_path, "r")
        total_transcription = f.read()
    
    len_total_transcription = len(total_transcription)

    MAX_INPUT_SIZE = 8000
    index = 0 

    total_translation = ""
    while len_total_transcription > index * MAX_INPUT_SIZE: 
        current_input = total_transcription[index * MAX_INPUT_SIZE : (index+1) * MAX_INPUT_SIZE]
        current_input_translation = translation_chinese_extraction(current_input)
        total_translation += current_input_translation

        index += 1

    return(total_translation)

#output#################################################################################################################
audio_path = "input.mp3" 
transcript_path = "transcription.txt"

print(transcription(audio_path) if option == 0 or option == 2 else "")
print(translation(transcript_path) if option == 1 or option == 2 else "")
